"""
app/services/parser.py
----------------------
Extracts plain text from uploaded resume files (PDF and DOCX).

Dependencies (see requirements.txt):
    pymupdf>=1.24.0     →  import fitz
    python-docx>=1.1.0  →  import docx
"""

from __future__ import annotations

import io
import re
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_text(file_path: str | Path) -> str:
    """
    Extract and clean plain text from a resume file.

    Supports:
        - .pdf  →  via PyMuPDF (fitz)
        - .docx →  via python-docx

    Args:
        file_path: Absolute or relative path to the uploaded resume file.

    Returns:
        Cleaned plain-text string. Empty string if extraction fails.

    Raises:
        ValueError: If the file extension is not supported.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        raw = _extract_pdf(path)
    elif ext == ".docx":
        raw = _extract_docx(path)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Only .pdf and .docx are accepted."
        )

    return _clean_text(raw)


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """
    Extract text directly from an in-memory byte buffer.
    Useful when you want to parse immediately after upload without
    writing to disk first.

    Args:
        content:  Raw file bytes.
        filename: Original filename (used only to detect extension).

    Returns:
        Cleaned plain-text string.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        raw = _extract_pdf_bytes(content)
    elif ext == ".docx":
        raw = _extract_docx_bytes(content)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Only .pdf and .docx are accepted."
        )

    return _clean_text(raw)


# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------

def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF file on disk using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.error("PyMuPDF is not installed. Run: pip install pymupdf")
        return ""

    try:
        doc = fitz.open(str(path))
        pages: list[str] = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return "\n".join(pages)
    except Exception as exc:
        logger.error("PDF extraction failed for '%s': %s", path, exc)
        return ""


def _extract_pdf_bytes(content: bytes) -> str:
    """Extract text from a PDF byte buffer using PyMuPDF."""
    try:
        import fitz
    except ImportError:
        logger.error("PyMuPDF is not installed. Run: pip install pymupdf")
        return ""

    try:
        doc = fitz.open(stream=content, filetype="pdf")
        pages: list[str] = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return "\n".join(pages)
    except Exception as exc:
        logger.error("PDF (bytes) extraction failed: %s", exc)
        return ""


# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------

def _extract_docx(path: Path) -> str:
    """Extract text from a DOCX file on disk using python-docx."""
    try:
        import docx
    except ImportError:
        logger.error("python-docx is not installed. Run: pip install python-docx")
        return ""

    try:
        document = docx.Document(str(path))
        paragraphs = [p.text for p in document.paragraphs]
        # Also capture text inside tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.error("DOCX extraction failed for '%s': %s", path, exc)
        return ""


def _extract_docx_bytes(content: bytes) -> str:
    """Extract text from a DOCX byte buffer using python-docx."""
    try:
        import docx
    except ImportError:
        logger.error("python-docx is not installed. Run: pip install python-docx")
        return ""

    try:
        document = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.error("DOCX (bytes) extraction failed: %s", exc)
        return ""


# ---------------------------------------------------------------------------
# Text cleaning
# ---------------------------------------------------------------------------

def _clean_text(text: str) -> str:
    """
    Normalise extracted text:
      1. Collapse 3+ consecutive blank lines → single blank line
      2. Strip leading/trailing whitespace per line
      3. Remove non-printable / control characters (except newline/tab)
      4. Strip overall leading/trailing whitespace
    """
    if not text:
        return ""

    # Remove control characters (keep \n and \t)
    text = re.sub(r"[^\x09\x0A\x20-\x7E\u00A0-\uFFFF]", " ", text)

    # Strip per-line trailing spaces
    lines = [line.rstrip() for line in text.splitlines()]

    # Collapse 3+ consecutive blank lines into one blank line
    cleaned: list[str] = []
    blank_count = 0
    for line in lines:
        if line.strip() == "":
            blank_count += 1
            if blank_count <= 1:
                cleaned.append("")
        else:
            blank_count = 0
            cleaned.append(line)

    return "\n".join(cleaned).strip()


# ---------------------------------------------------------------------------
# Utility: quick skill scanner (used by analyzer service)
# ---------------------------------------------------------------------------

# Common tech skills to scan for in the raw text
_SKILL_PATTERNS: list[str] = [
    "Python", "JavaScript", "TypeScript", "Java", "C\\+\\+", "C#", "Go", "Rust",
    "React", "Next\\.js", "Vue", "Angular", "Node\\.js", "Express",
    "FastAPI", "Django", "Flask", "Spring",
    "PostgreSQL", "MySQL", "MongoDB", "Redis", "SQLite",
    "Docker", "Kubernetes", "AWS", "GCP", "Azure",
    "GraphQL", "REST", "gRPC",
    "Git", "GitHub", "CI/CD", "Jenkins", "GitHub Actions",
    "HTML", "CSS", "Tailwind", "SASS",
    "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch", "scikit-learn",
    "Pandas", "NumPy", "SQL",
]

_SKILL_REGEX = re.compile(
    r"\b(" + "|".join(_SKILL_PATTERNS) + r")\b",
    flags=re.IGNORECASE,
)


def scan_skills(text: str) -> list[str]:
    """
    Return a deduplicated list of recognised skills found in the resume text.
    Preserves original casing of the first match.

    Args:
        text: Plain text extracted by extract_text() / extract_text_from_bytes().

    Returns:
        List of skill name strings, e.g. ['Python', 'React', 'Docker']
    """
    seen: dict[str, str] = {}  # lower → original-case
    for match in _SKILL_REGEX.finditer(text):
        key = match.group().lower()
        if key not in seen:
            seen[key] = match.group()
    return list(seen.values())
